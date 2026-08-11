"""Typesetting — real print output rendered from the locked PRESS style.

The renderers here turn the single manuscript source (``book_chapters`` via
PRESS) into distribution files:

  · a print-interior PDF with mirrored margins, KDP-scaled gutters, and the
    locked chapter-opening contract (header → title → approved epigraph);
  · a production DOCX from the same chapters under the same lock.

Both are deterministic given (style, chapters). The PDF render is the
authority on page count: ``actual_pages`` is whatever the renderer actually
produced, never a words-per-page estimate. The render also emits a
page → paragraph map so the EPUB page-list can reference REAL print pages.

Design rules:
  · the style lock drives everything — no font/size/trim decisions here;
  · a chapter with no text refuses the render (never typeset an empty book);
  · two-pass gutter sizing: the KDP inside margin depends on the page count,
    which depends on the margin — iterate until the bucket is stable.
"""

from __future__ import annotations

import re
from io import BytesIO
from typing import Any

from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import (
    BaseDocTemplate,
    Frame,
    NextPageTemplate,
    PageBreak,
    PageTemplate,
    Paragraph,
    Spacer,
)

from .atelier import TRIMS

__all__ = ["render_print_pdf", "render_docx", "TypesetError"]


class TypesetError(ValueError):
    """A render refusal with a user-readable reason."""


# ── KDP inside-margin (gutter) schedule: pages → required inches ─────────────
_GUTTER_SCHEDULE = [
    (150, 0.375),
    (300, 0.5),
    (500, 0.625),
    (700, 0.75),
    (828, 0.875),
]
_OUTER_MARGIN = 0.5
_TOP_MARGIN = 0.75
_BOTTOM_MARGIN = 0.75

# Base-14 PDF fonts — no embedding needed, valid in every reader.
_FONT_MAP = {
    "times": ("Times-Roman", "Times-Bold", "Times-Italic"),
    "garamond": ("Times-Roman", "Times-Bold", "Times-Italic"),
    "georgia": ("Times-Roman", "Times-Bold", "Times-Italic"),
    "helvetica": ("Helvetica", "Helvetica-Bold", "Helvetica-Oblique"),
    "arial": ("Helvetica", "Helvetica-Bold", "Helvetica-Oblique"),
    "courier": ("Courier", "Courier-Bold", "Courier-Oblique"),
}


def _gutter_for(pages: int) -> float:
    for limit, g in _GUTTER_SCHEDULE:
        if pages <= limit:
            return g
    return 1.0


def _font_family(name: str) -> tuple[str, str, str]:
    low = (name or "").lower()
    for key, fam in _FONT_MAP.items():
        if key in low:
            return fam
    return _FONT_MAP["times"]


def _pt(value: Any, default: float) -> float:
    """Parse '11', '11pt', 11, 11.0 → 11.0; refuse nonsense loudly."""
    if value is None or value == "":
        return default
    m = re.search(r"([\d.]+)", str(value))
    if not m:
        raise TypesetError(f"Unparseable size in style: {value!r}")
    return float(m.group(1))


def _check_inputs(style: dict, chapters: list[dict]) -> tuple[float, float]:
    trim = style.get("trim") or ""
    if trim not in TRIMS:
        raise TypesetError(f"Unknown trim '{trim}'. Valid: {', '.join(TRIMS)}")
    if not chapters:
        raise TypesetError("No chapters to typeset.")
    empty = [c["number"] for c in chapters if not (c.get("text") or "").strip()]
    if empty:
        raise TypesetError(f"Chapters without text cannot be typeset: {empty}")
    return TRIMS[trim]


def _paragraphs(text: str) -> list[str]:
    return [p.strip() for p in re.split(r"\n\s*\n", text) if p.strip()]


# ── PDF ──────────────────────────────────────────────────────────────────────


class _MirroredDoc(BaseDocTemplate):
    """Alternates odd/even page templates so the gutter always faces the
    spine, and records which (chapter, paragraph) starts each page."""

    def __init__(self, *args: Any, **kw: Any) -> None:
        super().__init__(*args, **kw)
        self.page_map: dict[int, tuple[int, int]] = {}
        self.chapter_pages: dict[int, int] = {}

    def handle_pageBegin(self) -> None:  # noqa: N802 (reportlab API)
        self._handle_pageBegin()
        # Page numbers are 1-based; page 1 is a recto (odd) page.
        nxt = "even" if (self.page + 1) % 2 == 0 else "odd"
        self._handle_nextPageTemplate(nxt)

    def afterFlowable(self, flowable: Any) -> None:  # noqa: N802
        anchor = getattr(flowable, "_ori_anchor", None)
        if anchor is not None and self.page not in self.page_map:
            self.page_map[self.page] = anchor
        ch_start = getattr(flowable, "_ori_chapter_start", None)
        if ch_start is not None and ch_start not in self.chapter_pages:
            self.chapter_pages[ch_start] = self.page


def _build_story(
    book: dict,
    chapters: list[dict],
    chapter_headers: dict[int, str],
    styles: dict[str, ParagraphStyle],
    frame_h: float,
) -> list[Any]:
    story: list[Any] = []
    if book.get("has_front"):
        story += [
            Spacer(1, frame_h * 0.35),
            Paragraph(_esc(book["title"]), styles["display"]),
            Spacer(1, 24),
            Paragraph(_esc(book["author_name"]), styles["center"]),
            NextPageTemplate("even"),
            PageBreak(),
            Spacer(1, frame_h * 0.7),
            Paragraph(
                _esc(f"Copyright © {book['author_name']}. All rights reserved."),
                styles["center"],
            ),
            PageBreak(),
        ]
    for ch in chapters:
        if story:
            story.append(PageBreak())
        opener = Spacer(1, frame_h * 0.28)
        opener._ori_chapter_start = ch["number"]  # type: ignore[attr-defined]
        story.append(opener)
        story.append(Paragraph(_esc(chapter_headers[ch["number"]]), styles["header"]))
        if ch.get("title"):
            story.append(Paragraph(_esc(ch["title"]), styles["title"]))
        epi = (ch.get("epigraph_text") or "").strip()
        if epi:
            story.append(Paragraph(_esc(epi).replace("\n", "<br/>"), styles["epigraph"]))
        for i, para in enumerate(_paragraphs(ch["text"])):
            p = Paragraph(_esc(para), styles["body_first"] if i == 0 else styles["body"])
            p._ori_anchor = (ch["number"], i)  # type: ignore[attr-defined]
            story.append(p)
    if book.get("has_back"):
        story += [
            PageBreak(),
            Spacer(1, frame_h * 0.35),
            Paragraph("About the Author", styles["title"]),
            Paragraph(_esc(book["author_name"]), styles["center"]),
        ]
    return story


def render_print_pdf(
    book: dict, style: dict, chapters: list[dict], chapter_headers: dict[int, str]
) -> dict:
    """Render the print interior. Returns bytes + the authoritative counts.

    ``chapters`` must carry full text (and ``epigraph_text`` when approved).
    ``chapter_headers`` maps chapter number → rendered header line
    ("Chapter Seven" / "Chapter 7" / "Chapter VII") from the locked style.
    """
    tw, th = _check_inputs(style, chapters)
    body_size = _pt(style.get("body_size"), 11.0)
    leading = _pt(style.get("leading"), round(body_size * 1.35, 1))
    if leading < body_size:
        raise TypesetError(f"Leading {leading} smaller than body size {body_size}.")
    body_font, body_bold, body_italic = _font_family(style.get("body_font", ""))
    head_font, head_bold, _ = _font_family(style.get("heading_font", ""))

    page_w, page_h = tw * inch, th * inch
    gutter = _gutter_for(300)  # first-pass guess; iterated below
    result: dict = {}

    for _ in range(4):  # gutter buckets converge fast; 4 passes is plenty
        buf = BytesIO()
        doc = _MirroredDoc(
            buf,
            pagesize=(page_w, page_h),
            title=book["title"],
            author=book["author_name"],
        )
        frame_w = page_w - (gutter + _OUTER_MARGIN) * inch
        frame_h = page_h - (_TOP_MARGIN + _BOTTOM_MARGIN) * inch
        odd = Frame(  # recto: gutter on the LEFT
            gutter * inch, _BOTTOM_MARGIN * inch, frame_w, frame_h, id="odd-frame"
        )
        even = Frame(  # verso: gutter on the RIGHT
            _OUTER_MARGIN * inch, _BOTTOM_MARGIN * inch, frame_w, frame_h, id="even-frame"
        )
        doc.addPageTemplates(
            [PageTemplate(id="odd", frames=[odd]), PageTemplate(id="even", frames=[even])]
        )

        st_body = ParagraphStyle(
            "body",
            fontName=body_font,
            fontSize=body_size,
            leading=leading,
            firstLineIndent=0.3 * inch,
            spaceAfter=0,
        )
        styles = {
            "body": st_body,
            "body_first": ParagraphStyle("body-first", parent=st_body, firstLineIndent=0),
            "header": ParagraphStyle(
                "ch-header",
                fontName=head_font,
                fontSize=body_size + 1,
                leading=(body_size + 1) * 1.2,
                alignment=1,
                spaceAfter=10,
            ),
            "title": ParagraphStyle(
                "ch-title",
                fontName=head_bold,
                fontSize=body_size + 5,
                leading=(body_size + 5) * 1.2,
                alignment=1,
                spaceAfter=18,
            ),
            "epigraph": ParagraphStyle(
                "epigraph",
                fontName=body_italic,
                fontSize=body_size - 1,
                leading=(body_size - 1) * 1.3,
                leftIndent=0.5 * inch,
                rightIndent=0.5 * inch,
                alignment=1,
                spaceAfter=18,
            ),
            "display": ParagraphStyle(
                "display",
                fontName=head_bold,
                fontSize=body_size + 9,
                leading=(body_size + 9) * 1.2,
                alignment=1,
            ),
            "center": ParagraphStyle(
                "center",
                fontName=body_font,
                fontSize=body_size,
                leading=leading,
                alignment=1,
            ),
        }
        story = _build_story(book, chapters, chapter_headers, styles, frame_h)
        doc.build(story)
        pages = doc.page
        result = {
            "pdf": buf.getvalue(),
            "actual_pages": pages,
            "gutter_in": gutter,
            "page_map": {p: doc.page_map[p] for p in sorted(doc.page_map)},
            "chapter_pages": dict(sorted(doc.chapter_pages.items())),
        }
        next_gutter = _gutter_for(pages)
        if next_gutter == gutter:
            break
        gutter = next_gutter
    return result


def _esc(s: str) -> str:
    return s.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


# ── DOCX ─────────────────────────────────────────────────────────────────────


def render_docx(
    book: dict, style: dict, chapters: list[dict], chapter_headers: dict[int, str]
) -> bytes:
    """Production DOCX from the same single source under the same lock."""
    _check_inputs(style, chapters)
    from docx import Document  # noqa: PLC0415 (heavy import, render-time only)
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK  # noqa: PLC0415
    from docx.oxml.ns import qn  # noqa: PLC0415
    from docx.shared import Inches, Pt  # noqa: PLC0415

    body_size = _pt(style.get("body_size"), 11.0)
    doc = Document()
    section = doc.sections[0]
    tw, th = TRIMS[style["trim"]]
    section.page_width, section.page_height = Inches(tw), Inches(th)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.5)
    section.top_margin = Inches(_TOP_MARGIN)
    section.bottom_margin = Inches(_BOTTOM_MARGIN)
    # Mirrored margins: left becomes the inside (gutter) edge on every spread.
    sect_pr = section._sectPr
    if sect_pr.find(qn("w:mirrorMargins")) is None:
        sect_pr.append(sect_pr.makeelement(qn("w:mirrorMargins"), {}))

    normal = doc.styles["Normal"]
    normal.font.name = style.get("body_font") or "Times New Roman"
    normal.font.size = Pt(body_size)

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run(book["title"])
    run.bold = True
    run.font.size = Pt(body_size + 9)
    by = doc.add_paragraph(book["author_name"])
    by.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for ch in chapters:
        doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
        h = doc.add_paragraph()
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        h.add_run(chapter_headers[ch["number"]]).font.name = (
            style.get("heading_font") or "Times New Roman"
        )
        if ch.get("title"):
            t = doc.add_paragraph()
            t.alignment = WD_ALIGN_PARAGRAPH.CENTER
            tr = t.add_run(ch["title"])
            tr.bold = True
            tr.font.size = Pt(body_size + 4)
        epi = (ch.get("epigraph_text") or "").strip()
        if epi:
            e = doc.add_paragraph()
            e.alignment = WD_ALIGN_PARAGRAPH.CENTER
            er = e.add_run(epi)
            er.italic = True
        for para in _paragraphs(ch["text"]):
            p = doc.add_paragraph(para)
            p.paragraph_format.first_line_indent = Inches(0.3)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
