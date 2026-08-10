"""Book export action.

Pulls all chapters from a Work's active book pipeline in sequence order
and assembles them into a single .docx manuscript file.
"""
from __future__ import annotations

import logging
from datetime import UTC, datetime
from pathlib import Path
from typing import TYPE_CHECKING

from orivellum.capabilities.actions import ActionBase

if TYPE_CHECKING:
    from orivellum.configuration.config import OrivellumConfig
    from orivellum.database.db import OrivellumDB

logger = logging.getLogger("orivellum.actions.book_export")


class BookExportAction(ActionBase):
    name = "book_export"
    description = (
        "Assemble a Work's book chapters in pipeline order into a single .docx "
        "manuscript file suitable for editing or submission."
    )
    category = "export"
    input_schema = {
        "type": "object",
        "properties": {
            "work_id": {"type": "string", "description": "The Work whose book to export"},
        },
        "required": ["work_id"],
    }

    def confirm_message(self, inputs: dict) -> str:
        return (
            "Assemble all extracted book chapters for this Work into a single "
            "**DOCX manuscript** in pipeline sequence order, ready to download."
        )

    def _execute_impl(self, inputs: dict, db: OrivellumDB, cfg: OrivellumConfig) -> dict:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Pt, RGBColor

        from orivellum.capabilities.generate import _now_label, _register_output, _slug

        work_id: str = inputs["work_id"]
        work = db.get_work(work_id)
        if not work:
            raise ValueError(f"Work {work_id!r} not found")

        # Get pipeline for this work
        pipeline = db.get_book_pipeline_for_work(work_id)
        if not pipeline:
            raise ValueError("No active book pipeline found for this Work. Start the book pipeline first.")

        pipeline_id = pipeline["id"]

        # Fetch all chapters for this pipeline ordered by seq
        with db._lock:
            rows = db._conn.execute(
                """SELECT seq, level, title, text, status, citations
                   FROM book_chapters
                   WHERE pipeline_id=?
                   ORDER BY seq""",
                (pipeline_id,),
            ).fetchall()

        chapters = [dict(r) for r in rows]
        if not chapters:
            raise ValueError("No chapters found in this pipeline. Run the extraction step first.")

        # ── Build .docx manuscript ──
        doc = Document()

        # Title page
        title_para = doc.add_heading(work.get("title", "Untitled Manuscript"), 0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if title_para.runs:
            title_para.runs[0].font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)

        desc = work.get("description") or ""
        if desc:
            p = doc.add_paragraph(desc)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if p.runs:
                p.runs[0].font.italic = True

        meta_p = doc.add_paragraph(
            f"Exported {datetime.now(UTC).strftime('%B %d, %Y')} "
            f"· {len(chapters)} chapter{'s' if len(chapters) != 1 else ''}"
        )
        meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if meta_p.runs:
            meta_p.runs[0].font.size = Pt(9)
            meta_p.runs[0].font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)
        doc.add_page_break()

        # Chapters
        for ch in chapters:
            level = max(1, min(int(ch.get("level") or 1), 4))
            ch_title = (ch.get("title") or f"Chapter {ch.get('seq', '')}").strip()
            doc.add_heading(ch_title, level)

            text = (ch.get("text") or "").strip()
            if text:
                # Split paragraphs on double-newline
                for para_text in text.split("\n\n"):
                    para_text = para_text.strip()
                    if para_text:
                        doc.add_paragraph(para_text)

            # Citations block (if present)
            citations_raw = ch.get("citations") or ""
            if citations_raw:
                import json as _json
                try:
                    cits = _json.loads(citations_raw) if isinstance(citations_raw, str) else citations_raw
                    if cits:
                        p = doc.add_paragraph()
                        r = p.add_run("Citations: " + "; ".join(str(c) for c in cits[:5]))
                        r.font.size = Pt(8)
                        r.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)
                except Exception:
                    pass

        # ── Save ──
        out_dir = Path(cfg.data_dir) / "outputs" / "generate" / work_id
        out_dir.mkdir(parents=True, exist_ok=True)
        slug = _slug(work.get("title", "book"))
        fname = f"{slug}_manuscript_{_now_label()}.docx"
        fpath = out_dir / fname
        doc.save(str(fpath))

        # Register output
        text_content = "\n\n".join(
            f"# {ch.get('title','')}\n{(ch.get('text') or '')[:500]}"
            for ch in chapters
        )
        title_out = f"Manuscript — {work.get('title', 'Book')}"
        doc_id = _register_output(fpath, work_id, db, cfg, "docx", title_out, text_content)

        data_dir = Path(cfg.data_dir)
        rel_path = str(fpath.relative_to(data_dir))

        return {
            "output_path": rel_path,
            "output_label": fname,
            "output_doc_id": doc_id,
            "chapter_count": len(chapters),
            "summary": f"Manuscript for '{work.get('title', 'Book')}' — {len(chapters)} chapters exported",
        }
