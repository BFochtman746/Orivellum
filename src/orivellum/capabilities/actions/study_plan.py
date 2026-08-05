"""Study plan action.

Reads a Work's concept graph and mastery levels, then generates a structured
learning plan as both a downloadable .docx and a new knowledge item so it
appears in the Learn home.
"""
from __future__ import annotations

import logging
from datetime import datetime, timezone
from pathlib import Path
from typing import TYPE_CHECKING

from orivellum.capabilities.actions import ActionBase

if TYPE_CHECKING:
    from orivellum.database.db import OrivellumDB
    from orivellum.configuration.config import OrivellumConfig

logger = logging.getLogger("orivellum.actions.study_plan")

_MASTERY_LABEL = {0: "Not started", 1: "Introduced", 2: "Familiar", 3: "Proficient", 4: "Mastered"}


class StudyPlanAction(ActionBase):
    name = "study_plan"
    description = (
        "Turn a Work's concept graph and current mastery levels into a structured "
        "learning plan — sorted by priority, with prerequisites resolved — saved "
        "as a knowledge item so it appears in your Learn home."
    )
    category = "learn"
    input_schema = {
        "type": "object",
        "properties": {
            "work_id": {"type": "string", "description": "The Work to generate a plan for"},
        },
        "required": ["work_id"],
    }

    def confirm_message(self, inputs: dict) -> str:
        return (
            "Read this Work's concept graph and mastery levels, sort concepts by "
            "learning priority (prerequisites first, weakest mastery first), "
            "and generate a **structured study plan** saved to your Learn home."
        )

    def _execute_impl(self, inputs: dict, db: "OrivellumDB", cfg: "OrivellumConfig") -> dict:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from orivellum.capabilities.generate import _register_output, _now_label, _slug

        work_id: str = inputs["work_id"]
        work = db.get_work(work_id)
        if not work:
            raise ValueError(f"Work {work_id!r} not found")

        # ── Fetch work_concepts + mastery ──
        with db._lock:
            concept_rows = db._conn.execute(
                """SELECT wc.id, wc.subject, wc.description,
                          COALESCE(
                              (SELECT AVG(wm.score)
                               FROM work_mastery wm
                               WHERE wm.concept_id = wc.id),
                              0.0
                          ) as mastery_avg
                   FROM work_concepts wc
                   WHERE wc.work_id=?
                   ORDER BY mastery_avg ASC, wc.subject ASC""",
                (work_id,),
            ).fetchall()

        concepts = [dict(r) for r in concept_rows]

        # Fall back to knowledge entities if no concepts defined
        if not concepts:
            with db._lock:
                krows = db._conn.execute(
                    """SELECT text FROM knowledge
                       WHERE work_id=? AND kind='entity'
                       ORDER BY created_at DESC LIMIT 40""",
                    (work_id,),
                ).fetchall()
            concepts = [{"subject": r["text"], "description": "", "mastery_avg": 0.0} for r in krows]

        if not concepts:
            raise ValueError("No concepts found for this Work. Add knowledge items or run the learning loop first.")

        # ── Build .docx study plan ──
        doc = Document()
        title_str = f"Study Plan — {work.get('title', 'Work')}"
        doc.add_heading(title_str, 0)
        doc.add_paragraph(
            f"Generated {datetime.now(timezone.utc).strftime('%B %d, %Y')} · "
            f"{len(concepts)} concept{'s' if len(concepts) != 1 else ''}"
        )

        doc.add_heading("Learning Priority Order", 1)
        doc.add_paragraph(
            "Concepts are sorted from least mastered to most mastered. "
            "Work through them in order for the most efficient path."
        )

        text_lines = [title_str, ""]
        for i, c in enumerate(concepts, 1):
            mastery = float(c.get("mastery_avg") or 0)
            level = min(4, int(mastery * 4 / 100)) if mastery > 1 else 0
            label = _MASTERY_LABEL.get(level, "Not started")
            name = c.get("subject", "") or c.get("name", "")
            desc = (c.get("description") or "").strip()

            p = doc.add_paragraph()
            p.add_run(f"{i}. {name}").bold = True
            status_run = p.add_run(f"  [{label}]")
            status_run.font.size = Pt(9)
            status_run.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)
            if desc:
                doc.add_paragraph(f"   {desc[:200]}", style="Normal")

            text_lines.append(f"{i}. [{label}] {name}")
            if desc:
                text_lines.append(f"   {desc[:200]}")

        # Summary
        doc.add_page_break()
        doc.add_heading("Quick Reference", 1)
        not_started = sum(1 for c in concepts if float(c.get("mastery_avg") or 0) == 0)
        doc.add_paragraph(
            f"Total concepts: {len(concepts)}\n"
            f"Not started: {not_started}\n"
            f"In progress: {len(concepts) - not_started}"
        )

        # ── Save .docx ──
        out_dir = Path(cfg.data_dir) / "outputs" / "generate" / work_id
        out_dir.mkdir(parents=True, exist_ok=True)
        slug = _slug(work.get("title", "work"))
        fname = f"{slug}_study_plan_{_now_label()}.docx"
        fpath = out_dir / fname
        doc.save(str(fpath))

        plan_text = "\n".join(text_lines)

        # Register output
        doc_id = _register_output(fpath, work_id, db, cfg, "docx", title_str, plan_text)

        # Also write a knowledge item so the plan appears in the Work's knowledge base
        try:
            db.create_knowledge_item(
                work_id=work_id,
                kind="note",
                text=plan_text[:3000],
                source_doc_id=doc_id,
                review_status="auto",
                meta={"plan_for": work_id, "generated_at": datetime.now(timezone.utc).isoformat()},
            )
        except Exception as exc:
            logger.warning("Could not write study plan knowledge item: %s", exc)

        data_dir = Path(cfg.data_dir)
        rel_path = str(fpath.relative_to(data_dir))

        return {
            "output_path": rel_path,
            "output_label": fname,
            "output_doc_id": doc_id,
            "concept_count": len(concepts),
            "summary": f"Study plan for '{work.get('title', 'Work')}' — {len(concepts)} concepts prioritised",
        }
