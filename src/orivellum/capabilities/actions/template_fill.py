"""Template fill action.

Accepts a previously uploaded .docx or .xlsx template (by content_path or doc_id)
and a data dict, renders the template, and returns the filled file.

.docx  — rendered via docxtpl (Jinja2 inside DOCX)
.xlsx  — rendered via openpyxl (named-range / cell substitution)
"""

from __future__ import annotations

import json
import logging
from pathlib import Path
from typing import TYPE_CHECKING

from orivellum.capabilities.actions import ActionBase

if TYPE_CHECKING:
    from orivellum.configuration.config import OrivellumConfig
    from orivellum.database.db import OrivellumDB

logger = logging.getLogger("orivellum.actions.template_fill")


class TemplateFillAction(ActionBase):
    name = "template_fill"
    description = (
        "Fill a .docx or .xlsx template with structured data. "
        "The template uses {{variable}} placeholders (docx) or named cells (xlsx). "
        "Provide a document ID or file path plus a JSON data object."
    )
    category = "generate"
    input_schema = {
        "type": "object",
        "properties": {
            "template_doc_id": {
                "type": "string",
                "description": "Library document ID of the template file",
            },
            "data": {
                "type": "object",
                "description": "Key-value pairs to substitute into the template",
            },
            "output_name": {
                "type": "string",
                "description": "Optional filename for the output (without extension)",
            },
            "work_id": {
                "type": "string",
                "description": "Optional Work to associate the output with",
            },
        },
        "required": ["template_doc_id", "data"],
    }

    def confirm_message(self, inputs: dict) -> str:
        data = inputs.get("data", {})
        field_count = len(data) if isinstance(data, dict) else "?"
        return (
            f"Fill the selected template with **{field_count} data field{'s' if field_count != 1 else ''}** "
            f"and produce a ready-to-download filled copy."
        )

    def _execute_impl(self, inputs: dict, db: OrivellumDB, cfg: OrivellumConfig) -> dict:
        from orivellum.capabilities.generate import _now_label, _register_output

        template_doc_id: str = inputs["template_doc_id"]
        data: dict = inputs.get("data", {})
        work_id: str | None = inputs.get("work_id")
        output_name: str | None = inputs.get("output_name")

        if isinstance(data, str):
            data = json.loads(data)

        # Look up template document
        template_doc = db.get_document(template_doc_id)
        if not template_doc:
            raise ValueError(f"Template document {template_doc_id!r} not found in library")

        cp = template_doc.get("content_path")
        if not cp:
            raise ValueError("Template document has no stored file (content_path is empty)")

        data_dir = Path(cfg.data_dir)
        template_path = data_dir / cp
        if not template_path.is_file():
            raise ValueError(f"Template file not found on disk: {cp}")

        suffix = template_path.suffix.lower()
        ts = _now_label()
        out_name = f"{output_name or 'filled'}_{ts}{suffix}"
        out_dir = data_dir / "outputs" / "generate" / (work_id or "library")
        out_dir.mkdir(parents=True, exist_ok=True)
        out_path = out_dir / out_name

        text_content = json.dumps(data)[:2000]

        if suffix == ".docx":
            out_path, text_content = _fill_docx(template_path, out_path, data)
        elif suffix in (".xlsx", ".xls"):
            out_path, text_content = _fill_xlsx(template_path, out_path, data)
        else:
            raise ValueError(f"Unsupported template type: {suffix} (must be .docx or .xlsx)")

        # Register output
        doc_id = _register_output(
            out_path,
            work_id or None,
            db,
            cfg,
            suffix.lstrip("."),
            output_name or f"Filled {template_doc.get('title', 'Template')}",
            text_content,
        )

        rel_path = str(out_path.relative_to(data_dir))
        return {
            "output_path": rel_path,
            "output_label": out_name,
            "output_doc_id": doc_id,
            "summary": f"Template filled with {len(data)} fields → {out_name}",
        }


def _fill_docx(template_path: Path, out_path: Path, data: dict) -> tuple[Path, str]:
    """Render a .docx template via docxtpl (Jinja2) with fallback to regex replace."""
    try:
        from docxtpl import DocxTemplate

        tpl = DocxTemplate(str(template_path))
        tpl.render(data)
        tpl.save(str(out_path))
        text_content = " ".join(str(v) for v in data.values())
        return out_path, text_content
    except ImportError:
        # Fallback: plain text replacement via python-docx
        from docx import Document

        doc = Document(str(template_path))
        for para in doc.paragraphs:
            for key, val in data.items():
                placeholder = f"{{{{{key}}}}}"
                if placeholder in para.text:
                    for run in para.runs:
                        run.text = run.text.replace(placeholder, str(val))
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for key, val in data.items():
                        placeholder = f"{{{{{key}}}}}"
                        if placeholder in cell.text:
                            for para in cell.paragraphs:
                                for run in para.runs:
                                    run.text = run.text.replace(placeholder, str(val))
        doc.save(str(out_path))
        text_content = " ".join(str(v) for v in data.values())
        return out_path, text_content


def _fill_xlsx(template_path: Path, out_path: Path, data: dict) -> tuple[Path, str]:
    """Fill an .xlsx template by replacing {{key}} in all cells."""
    import openpyxl

    wb = openpyxl.load_workbook(str(template_path))
    for ws in wb.worksheets:
        for row in ws.iter_rows():
            for cell in row:
                if cell.data_type == "s" and cell.value:
                    val = str(cell.value)
                    for key, replacement in data.items():
                        val = val.replace(f"{{{{{key}}}}}", str(replacement))
                    cell.value = val
    wb.save(str(out_path))
    text_content = " ".join(str(v) for v in data.values())
    return out_path, text_content
