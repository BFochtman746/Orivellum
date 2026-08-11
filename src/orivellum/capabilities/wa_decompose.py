"""Writing Architect archive decomposition — Pipeline M0 (DECOMPOSE).

Converts the WRITING_ARCHITECT archive (engine doctrine, runtime policies,
story-system schemas, trilogy bible data, voice material) from prose
documents into machine-readable records:

- Coverage inventory: every real file hashed, deduped (``__N`` /
  trailing-" N" duplicate pairs), and given an explicit disposition —
  ``extracted``, ``deduped``, or ``deferred`` with a reason.  Nothing is
  silently dropped.
- ENGINE docs → Engine Contract records (purpose, inputs, allowed /
  forbidden operations, checks, thresholds) with certification status
  backfilled from ENGINE_INDEX.
- RUNTIME docs → runtime policy / record schema entries;
  RELIABILITY_LAYER docs → reliability policies.
- STORY_SYSTEM schemas → table specs; other story docs → story policies.
- BIBLE_DATA + Story Bible + Book Bible → canon-fact PROPOSALS only
  (HISTORICAL / INFERRED / INVENTED), scoped to the trilogy, queued for
  author ratification.  The decomposer never writes canon authority.
- Voice Architect / Held-Breath / Ultimate Prose / Prose Style Reference →
  one measurable voice-envelope spec; Ultimate Diagnostic → POSITION audit
  spec; AI Provenance v2.0 + Module 7 → provenance spec.
"""

from __future__ import annotations

import hashlib
import io
import json
import logging
import re
import time
import uuid
import zipfile
from dataclasses import dataclass, field
from pathlib import Path

logger = logging.getLogger("orivellum.wa.decompose")

# Scope every canon proposal to the trilogy the BIBLE_DATA declares.
TRILOGY_SCOPE = "series:The Harp in the Dark"

_SKIP_PARTS = ("__MACOSX",)
_MAX_FACT_CHARS = 2400
_MIN_FACT_CHARS = 40

# Zip-bomb / runaway-archive ceilings (the real archive is ~200 files, <100 MB)
_MAX_ENTRIES = 5000
_MAX_MEMBER_BYTES = 100 * 1024 * 1024
_MAX_TOTAL_BYTES = 1024 * 1024 * 1024

# ── Section parsing ───────────────────────────────────────────────────────────


@dataclass
class Section:
    heading: str
    level: int
    items: list[str] = field(default_factory=list)


@dataclass
class ParsedDoc:
    title: str
    sections: list[Section] = field(default_factory=list)
    tables: list[list[list[str]]] = field(default_factory=list)
    frontmatter: dict = field(default_factory=dict)

    def section_map(self) -> dict[str, list[str]]:
        out: dict[str, list[str]] = {}
        for s in self.sections:
            key = s.heading.strip().lower()
            out.setdefault(key, []).extend(s.items)
        return out

    def all_text(self) -> str:
        parts: list[str] = []
        for s in self.sections:
            if s.heading:
                parts.append(s.heading)
            parts.extend(s.items)
        return "\n".join(parts)


def parse_docx_bytes(data: bytes, fallback_title: str) -> ParsedDoc:
    """Heading-aware parse of a .docx: sections keyed by Heading styles."""
    import docx  # python-docx — already a project dependency

    d = docx.Document(io.BytesIO(data))
    doc = ParsedDoc(title=fallback_title)
    current = Section(heading="", level=0)
    doc.sections.append(current)
    for para in d.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style = (para.style.name or "") if para.style is not None else ""
        m = re.match(r"heading\s*(\d+)", style, re.IGNORECASE)
        if m or style.lower() == "title":
            level = int(m.group(1)) if m else 0
            current = Section(heading=text, level=level)
            doc.sections.append(current)
        else:
            prefix = "- " if "list" in style.lower() else ""
            current.items.append(prefix + text)
    for table in d.tables:
        rows = []
        for tr in table.rows:
            rows.append([cell.text.strip() for cell in tr.cells])
        if rows:
            doc.tables.append(rows)
    # First non-heading line is usually the document's own name
    if doc.sections and doc.sections[0].items:
        doc.title = doc.sections[0].items[0].lstrip("- ").strip() or fallback_title
    return doc


def parse_markdown(text: str, fallback_title: str) -> ParsedDoc:
    """Parse markdown with optional YAML frontmatter into heading sections."""
    doc = ParsedDoc(title=fallback_title)
    body = text
    if text.startswith("---"):
        end = text.find("\n---", 3)
        if end != -1:
            fm_block = text[3:end]
            body = text[end + 4 :]
            for line in fm_block.splitlines():
                m = re.match(r"^(\w[\w_-]*):\s*(.*)$", line.strip())
                if m:
                    doc.frontmatter[m.group(1)] = m.group(2).strip("'\" ")
    current = Section(heading="", level=0)
    doc.sections.append(current)
    for line in body.splitlines():
        stripped = line.strip()
        if not stripped:
            continue
        m = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if m:
            current = Section(heading=m.group(2).strip(), level=len(m.group(1)))
            doc.sections.append(current)
        else:
            current.items.append(stripped)
    if doc.frontmatter.get("title"):
        doc.title = doc.frontmatter["title"]
    return doc


# ── Path classification ───────────────────────────────────────────────────────


def layer_of(rel_path: str) -> str:
    p = rel_path
    if "/WRITING_SYSTEM/ENGINE/" in p:
        return "ENGINE"
    if "/WRITING_SYSTEM/RUNTIME/" in p:
        return "RUNTIME"
    if "/WRITING_SYSTEM/RELIABILITY_LAYER/" in p:
        return "RELIABILITY"
    if "/WRITING_SYSTEM/STORY_SYSTEM/" in p:
        return "STORY_SYSTEM"
    if "/WRITING_SYSTEM/BIBLE_DATA/" in p:
        return "BIBLE_DATA"
    if "/WRITING_SYSTEM/" in p:
        return "WRITING_SYSTEM"
    if "/UMS_v23" in p:
        return "UMS"
    if "UNHINDERED_MASTERY_SYSTEM" in p:
        return "UNHINDERED"
    if "/Module writing/" in p:
        return "MODULE_WRITING"
    return "ROOT"


# Root files with special handling. Everything not matched below is deferred
# with an explicit reason (never silently dropped).
_VOICE_SOURCES = {
    "The_Voice_Architect.docx",
    "The_Held-Breath_Voice.docx",
    "ULTIMATE_PROSE_SYSTEM_v2.0_COMPLETE.docx",
    "ULTIMATE_PROSE_SYSTEM_v2.1_IMPLEMENTATION.md",
    "01_-_Prose_Style_Reference.md",
}
_DEFER_REASONS: list[tuple[re.Pattern, str]] = [
    (re.compile(r"NARRATIVEOS_v(21|24_1|24_3)"), "Superseded by NARRATIVEOS v24.4"),
    (
        re.compile(r"NARRATIVEOS_v24_4"),
        "Master-system prose — consumed by later engine milestones (LOOM/FORCE)",
    ),
    (
        re.compile(r"AI_Provenance_Verification_System_v1\.0"),
        "Superseded by AI Provenance v2.0",
    ),
    (
        re.compile(r"FORGE_v3_3"),
        "FORGE dimensions already integrated natively; Tier-3 registration in ASSAY milestone",
    ),
    (
        re.compile(r"Unified_Writers_Room|UNHINDERED_Mastery_Operating_System"),
        "Writers-room role doctrine — consumed by the LOOM drafting milestone",
    ),
    (
        re.compile(r"SOVEREIGN"),
        "SOVEREIGN master/system material — superseded by the Sovereign Platform build",
    ),
    (
        re.compile(
            r"Ultimate_Biblical_Research|Ultimately_Character_Profile|Ultimately_World_Build"
        ),
        "Research/authoring prompt template — consumed by research milestones",
    ),
    (
        re.compile(r"Complete_Author_Profile|Universal_Reader_Copy"),
        "Author/reader profile material — consumed by voice & distribution milestones",
    ),
    (
        re.compile(r"ULTIMATE_AI_DETECTION"),
        "Detection implementation plan — instrument registration in ASSAY milestone",
    ),
    (
        re.compile(r"Writing_Architect_Complete_Audit_Report"),
        "Audit narrative about the archive itself — no doctrine to extract",
    ),
    (re.compile(r"CFNetworkDownload"), "Browser download artifact — no doctrine content"),
    (
        re.compile(r"CH\d+_BLUEPRINT|MANUSCRIPT_MASTER"),
        "Manuscript/blueprint content — consumed by later milestones (LOOM/BAND), not doctrine",
    ),
    (
        re.compile(r"00_-_Master_Command_Center|07_-_Dataview|08_-_Templater|09_-_Graph_View"),
        "Obsidian vault tooling configuration — no doctrine to extract",
    ),
]

# BIBLE_DATA classification defaults; per-item [HISTORICAL]/[INFERRED]/[INVENTED]
# markers in the text override these.
_BIBLE_CLASSIFICATION: dict[str, str] = {
    "Iron_Age_Material_Culture.md": "HISTORICAL",
    "Hebrew_Name_Etymology_Bank.md": "HISTORICAL",
    "Kingship_Ideology.md": "HISTORICAL",
    "God_as_Character_UPDATED.md": "INFERRED",
    "Index_-_Motifs_-_Volume_II.md": "INVENTED",
}
_CLASS_MARKER = re.compile(r"\[(HISTORICAL|INFERRED|INVENTED)\]", re.IGNORECASE)


def defer_reason_for(rel_path: str, filename: str) -> str | None:
    for pat, reason in _DEFER_REASONS:
        if pat.search(filename) or pat.search(rel_path):
            return reason
    if filename.lower().endswith(".zip"):
        return "Nested archive — extraction deferred to a later milestone"
    layer = layer_of(rel_path)
    if layer in ("UMS", "UNHINDERED", "MODULE_WRITING"):
        return "Legacy mastery-system material — superseded by the WRITING_SYSTEM doctrine"
    return None


# ── Engine contracts ──────────────────────────────────────────────────────────

_CONTRACT_KEYS = {
    "purpose": ("purpose",),
    "required_inputs": ("required inputs", "inputs"),
    "allowed_operations": ("allowed operations",),
    "forbidden_operations": ("forbidden operations",),
    "authority_relationship": ("authority relationship",),
    "output_schema": ("required output schema", "output schema", "outputs"),
    "execution_steps": ("execution steps",),
    "stop_conditions": ("stop conditions",),
    "failure_states": ("failure states",),
    "escalation_path": ("escalation path",),
    "verification": ("verification method", "verification"),
    "stress_tests": ("stress tests",),
}


def build_engine_contract(name: str, doc: ParsedDoc) -> dict:
    smap = doc.section_map()
    contract: dict = {"role": "engine" if re.match(r"ENGINE_\d+", name) else "support"}
    for key, aliases in _CONTRACT_KEYS.items():
        for alias in aliases:
            if alias in smap and smap[alias]:
                contract[key] = [i.lstrip("- ").strip() for i in smap[alias]]
                break
    status = _CLASS_STATUS.search(doc.all_text())
    if status:
        contract["certification_status"] = status.group(1).upper()
    contract["sections"] = [
        {"heading": s.heading, "items": s.items} for s in doc.sections if s.heading or s.items
    ]
    if doc.tables:
        contract["tables"] = doc.tables
    return contract


_CLASS_STATUS = re.compile(r"\b(PRESERVED|PARTIAL|CERTIFIED|DEPRECATED)\b")


def engine_index_statuses(doc: ParsedDoc) -> dict[str, dict[str, str]]:
    """Extract engine → index metadata from the ENGINE_INDEX operator table.

    The index table columns are File / Purpose / When to Call / Runtime
    Status (e.g. "Runtime Required", "Runtime Core", "On-Demand").
    """
    out: dict[str, dict[str, str]] = {}
    for table in doc.tables:
        if not table:
            continue
        header = [h.strip().lower() for h in table[0]]
        try:
            fi = header.index("file")
            si = header.index("runtime status")
        except ValueError:
            continue
        pi = header.index("purpose") if "purpose" in header else None
        wi = header.index("when to call") if "when to call" in header else None
        for row in table[1:]:
            if len(row) <= max(fi, si) or not row[fi].strip():
                continue
            key = re.sub(r"\W+", "_", row[fi].strip().upper()).strip("_")
            meta = {"certification_status": row[si].strip()}
            if pi is not None and len(row) > pi and row[pi].strip():
                meta["index_purpose"] = row[pi].strip()
            if wi is not None and len(row) > wi and row[wi].strip():
                meta["when_to_call"] = row[wi].strip()
            out[key] = meta
    return out


# ── Canon proposals ───────────────────────────────────────────────────────────


def proposals_from_doc(doc: ParsedDoc, rel_path: str, default_classification: str) -> list[dict]:
    """One proposal per heading section, with per-item marker overrides."""
    out: list[dict] = []
    for s in doc.sections:
        if not s.items:
            continue
        text = "\n".join(s.items)
        if len(text) < _MIN_FACT_CHARS:
            continue
        text = text[:_MAX_FACT_CHARS]
        marker = _CLASS_MARKER.search(s.heading + " " + text[:200])
        classification = marker.group(1).upper() if marker else default_classification
        title = s.heading or doc.title
        location = s.heading or "(preamble)"
        pid = hashlib.sha256(f"{rel_path}|{location}|{text}".encode()).hexdigest()[:32]
        out.append(
            {
                "id": pid,
                "fact_title": title[:300],
                "fact_text": text,
                "classification": classification,
                "scope": TRILOGY_SCOPE,
                "source_path": rel_path,
                "source_location": location[:300],
            }
        )
    return out


# ── Voice envelope / POSITION / provenance specs ──────────────────────────────

_MEASURABLE = re.compile(
    r"\d|%|\bnever\b|\balways\b|\bmust(?:\s+not)?\b|\bno more than\b|\bat least\b|\bban\b|\bforbid",
    re.IGNORECASE,
)


def distill_voice_envelope(sources: dict[str, ParsedDoc]) -> dict:
    """Distill voice doctrine into a measurable envelope spec.

    Keeps only rules that are checkable (contain a number, percentage, or a
    hard modal like never/always/must), grouped by source, plus the full
    section outline for traceability.
    """
    constraints: dict[str, list[dict]] = {}
    outline: dict[str, list[str]] = {}
    for name, doc in sources.items():
        rules: list[dict] = []
        for s in doc.sections:
            for item in s.items:
                clean = item.lstrip("- ").strip()
                if len(clean) >= 15 and _MEASURABLE.search(clean):
                    rules.append({"rule": clean[:400], "section": s.heading[:120]})
        constraints[name] = rules[:200]
        outline[name] = [s.heading for s in doc.sections if s.heading]
    return {
        "scope": TRILOGY_SCOPE,
        "precedence": (
            "Prose Style Reference is the series-scoped sentence-level authority; "
            "it supersedes general voice instructions when in conflict."
        ),
        "measurable_constraints": constraints,
        "source_outlines": outline,
    }


def build_position_spec(doc: ParsedDoc) -> dict:
    """Map the Ultimate Diagnostic into POSITION audit dimensions."""
    dimensions: list[dict] = []
    current: dict | None = None
    for s in doc.sections:
        for item in s.items:
            clean = item.lstrip("- ").strip()
            m = re.match(r"^(\d+)\.\s*(.+)$", clean)
            if m:
                current = {"index": int(m.group(1)), "dimension": m.group(2), "probes": []}
                dimensions.append(current)
            elif current is not None and len(clean) > 10:
                current["probes"].append(clean[:400])
    return {
        "dimensions": dimensions,
        "rule": "Derive manuscript stage from evidence, never from claimed status.",
        "raw_sections": [{"heading": s.heading, "items": s.items} for s in doc.sections],
    }


def build_generic_payload(doc: ParsedDoc) -> dict:
    payload: dict = {
        "title": doc.title,
        "sections": [
            {"heading": s.heading, "items": s.items} for s in doc.sections if s.heading or s.items
        ],
    }
    if doc.tables:
        payload["tables"] = doc.tables
    if doc.frontmatter:
        payload["frontmatter"] = doc.frontmatter
    return payload


# ── Orchestrator ──────────────────────────────────────────────────────────────


def _dedupe_key(rel_path: str) -> str:
    """Group ``NAME__1.docx`` / ``NAME 2`` variants with their base file."""
    parent, _, name = rel_path.rpartition("/")
    stem, dot, ext = name.rpartition(".")
    # Only treat the suffix as a real extension (e.g. ".docx"), not a
    # version fragment like the ".4 2" in "SOVEREIGN_MASTER_v1.4 2".
    if not dot or not re.fullmatch(r"[A-Za-z0-9]{1,6}", ext):
        stem, ext = name, ""
    base = re.sub(r"__\d+$", "", stem)
    base = re.sub(r"\s+\d+$", "", base)
    return f"{parent}/{base}.{ext}" if ext else f"{parent}/{base}"


def _resolve_duplicates(
    groups: dict[str, list[str]],
    hashes: dict[str, str],
    sizes: dict[str, int],
) -> dict[str, tuple[str, str | None, str | None]]:
    """Resolve each duplicate group to one canonical file.

    Canonical = the most complete revision: largest file wins, base
    filename breaks ties. Every other member is resolved explicitly
    (deduped if byte-identical, deferred with a reason if it differs).
    """
    dupe_status: dict[str, tuple[str, str | None, str | None]] = {}
    for key, members in groups.items():
        members_sorted = sorted(members, key=lambda p: (-sizes[p], p != key, p))
        canon = members_sorted[0]
        for other in members_sorted[1:]:
            if hashes[other] == hashes[canon]:
                dupe_status[other] = ("deduped", canon, "Byte-identical duplicate")
            else:
                dupe_status[other] = (
                    "deferred",
                    canon,
                    "Variant differs from canonical copy (hash mismatch) — "
                    "needs manual reconciliation",
                )
    return dupe_status


def _finalize_engine_records(
    engine_docs: list[tuple[str, str, ParsedDoc]], records: list[dict]
) -> None:
    """Build engine contracts, backfilling metadata (runtime status,
    purpose, when-to-call) from the ENGINE_INDEX operator table."""
    index_meta: dict[str, dict[str, str]] = {}
    for _, filename, doc in engine_docs:
        if filename.startswith("ENGINE_INDEX"):
            index_meta.update(engine_index_statuses(doc))
    for rel, filename, doc in engine_docs:
        name = Path(filename).stem
        payload = build_engine_contract(name, doc)
        key = re.sub(r"\W+", "_", name.upper()).strip("_")
        meta = index_meta.get(key) or next(
            (m for k, m in index_meta.items() if k in key or key in k), None
        )
        if meta:
            for mk, mv in meta.items():
                payload.setdefault(mk, mv)
        records.append(_record("engine_contract", name, payload, rel))


def _safe_entries(zf: zipfile.ZipFile) -> list[zipfile.ZipInfo]:
    """List real archive members, enforcing zip-bomb ceilings."""
    entries = [
        info
        for info in zf.infolist()
        if not info.is_dir()
        and not any(part in info.filename.split("/") for part in _SKIP_PARTS)
        and not Path(info.filename).name.startswith("._")
    ]
    if len(entries) > _MAX_ENTRIES:
        raise ValueError(f"Archive has {len(entries)} entries (max {_MAX_ENTRIES})")
    total_bytes = sum(info.file_size for info in entries)
    if total_bytes > _MAX_TOTAL_BYTES:
        raise ValueError(f"Archive expands to {total_bytes} bytes (max {_MAX_TOTAL_BYTES})")
    oversized = [i.filename for i in entries if i.file_size > _MAX_MEMBER_BYTES]
    if oversized:
        raise ValueError(f"Archive member too large: {oversized[0]}")
    return entries


def run_decompose(archive_path: Path, db, data_dir: Path | None = None) -> dict:
    """Decompose the WRITING_ARCHITECT archive into machine-readable records.

    Reads the zip directly, builds the inventory with dedupe resolution,
    extracts doctrine records and canon proposals, persists everything via
    WAStore, and writes a coverage report proving every file is accounted
    for. Returns the summary dict.
    """
    from orivellum.database.wa_store import WAStore

    t0 = time.monotonic()
    run_id = str(uuid.uuid4())
    inventory: list[dict] = []
    records: list[dict] = []
    proposals: list[dict] = []
    voice_sources: dict[str, ParsedDoc] = {}
    engine_docs: list[tuple[str, str, ParsedDoc]] = []  # (rel_path, filename, doc)

    with zipfile.ZipFile(archive_path) as zf:
        entries = _safe_entries(zf)
        # Pass 1 — hash everything and resolve duplicates
        hashes: dict[str, str] = {}
        groups: dict[str, list[str]] = {}
        for info in entries:
            data = zf.read(info)
            hashes[info.filename] = hashlib.sha256(data).hexdigest()
            groups.setdefault(_dedupe_key(info.filename), []).append(info.filename)

        sizes = {info.filename: info.file_size for info in entries}
        dupe_status = _resolve_duplicates(groups, hashes, sizes)

        # Pass 2 — extract canonical files per layer
        for info in entries:
            rel = info.filename
            filename = Path(rel).name
            layer = layer_of(rel)
            row = {
                "id": str(uuid.uuid4()),
                "rel_path": rel,
                "filename": filename,
                "layer": layer,
                "sha256": hashes[rel],
                "size_bytes": info.file_size,
                "duplicate_of": None,
                "status": "extracted",
                "reason": None,
                "target_kind": None,
            }
            if rel in dupe_status:
                status, canon, reason = dupe_status[rel]
                row.update(status=status, duplicate_of=canon, reason=reason)
                inventory.append(row)
                continue

            reason = defer_reason_for(rel, filename)
            if reason is not None:
                row.update(status="deferred", reason=reason)
                inventory.append(row)
                continue

            data = zf.read(info)
            try:
                doc = _parse_any(data, filename)
            except Exception as exc:
                logger.warning("Parse failed for %s: %s", rel, exc)
                row.update(status="deferred", reason=f"Unparseable ({exc})")
                inventory.append(row)
                continue
            if doc is None:
                row.update(status="deferred", reason="Unsupported file type for extraction")
                inventory.append(row)
                continue

            target = _route_document(
                rel, filename, layer, doc, records, proposals, voice_sources, engine_docs
            )
            row["target_kind"] = target
            inventory.append(row)

    _finalize_engine_records(engine_docs, records)

    # Voice envelope — one distilled spec from all voice sources
    if voice_sources:
        records.append(
            _record(
                "voice_envelope",
                "voice_envelope_v1",
                distill_voice_envelope(voice_sources),
                "; ".join(sorted(voice_sources)),
                note=(
                    "Distilled from Voice Architect, Held-Breath, "
                    "Ultimate Prose, Prose Style Reference"
                ),
            )
        )

    store = WAStore(db)
    counts = store.replace_run(run_id, inventory, records, proposals)
    coverage = store.coverage()
    summary = {
        "run_id": run_id,
        "archive": str(archive_path),
        "duration_sec": round(time.monotonic() - t0, 2),
        **counts,
        "coverage": coverage,
    }
    if data_dir is not None:
        _write_coverage_report(data_dir, summary, inventory)
    logger.info(
        "WA decompose: %d docs (%s), %d records, %d proposals in %.1fs",
        counts["inventory"],
        coverage.get("by_status"),
        counts["records"],
        counts["proposals_seen"],
        summary["duration_sec"],
    )
    return summary


def _parse_any(data: bytes, filename: str) -> ParsedDoc | None:
    lower = filename.lower()
    stem = Path(filename).stem
    if lower.endswith(".docx"):
        return parse_docx_bytes(data, stem)
    if lower.endswith((".md", ".txt")):
        return parse_markdown(data.decode("utf-8", errors="replace"), stem)
    if lower.endswith(".json"):
        doc = ParsedDoc(title=stem)
        try:
            payload = json.loads(data.decode("utf-8", errors="replace"))
        except Exception:
            return None
        sec = Section(heading="json", level=1)
        sec.items.append(json.dumps(payload, ensure_ascii=False)[:20000])
        doc.sections.append(sec)
        return doc
    return None


def _record(
    record_type: str, name: str, payload: dict, source: str, note: str | None = None
) -> dict:
    return {
        "id": str(uuid.uuid4()),
        "record_type": record_type,
        "name": name,
        "payload": payload,
        "source_path": source,
        "source_note": note,
    }


def _route_document(
    rel: str,
    filename: str,
    layer: str,
    doc: ParsedDoc,
    records: list[dict],
    proposals: list[dict],
    voice_sources: dict[str, ParsedDoc],
    engine_docs: list[tuple[str, str, ParsedDoc]],
) -> str:
    """Route one parsed canonical document to its extraction target.

    Returns the target_kind stored in the inventory row.
    """
    if filename in _VOICE_SOURCES:
        voice_sources[filename] = doc
        return "voice_envelope"
    if layer == "ENGINE":
        engine_docs.append((rel, filename, doc))
        return "engine_contract"
    handled = _route_layer_doc(rel, filename, layer, doc, records, proposals)
    if handled is not None:
        return handled
    return _route_root_doc(rel, filename, doc, records, proposals)


def _route_layer_doc(
    rel: str,
    filename: str,
    layer: str,
    doc: ParsedDoc,
    records: list[dict],
    proposals: list[dict],
) -> str | None:
    """Route WRITING_SYSTEM layer docs; returns None for root-level files."""
    stem = Path(filename).stem
    if layer == "RUNTIME":
        rtype = "record_schema" if filename.startswith("schema_") else "runtime_policy"
        records.append(_record(rtype, stem, build_generic_payload(doc), rel))
        return rtype
    if layer == "RELIABILITY":
        records.append(_record("reliability_policy", stem, build_generic_payload(doc), rel))
        return "reliability_policy"
    if layer == "STORY_SYSTEM":
        if filename.startswith(("STORY_BIBLE_MASTER", "STORY_BIBLE_ACTIVE")):
            proposals.extend(proposals_from_doc(doc, rel, "INVENTED"))
            return "canon_proposal"
        rtype = "table_spec" if "_SCHEMA" in filename.upper() else "story_policy"
        records.append(_record(rtype, stem, build_generic_payload(doc), rel))
        return rtype
    if layer == "BIBLE_DATA":
        if filename == "06_-_Scene_Pressure_Matrix.md":
            records.append(_record("pressure_matrix", stem, build_generic_payload(doc), rel))
            return "pressure_matrix"
        default = _BIBLE_CLASSIFICATION.get(filename, "INFERRED")
        proposals.extend(proposals_from_doc(doc, rel, default))
        return "canon_proposal"
    return None


def _route_root_doc(
    rel: str,
    filename: str,
    doc: ParsedDoc,
    records: list[dict],
    proposals: list[dict],
) -> str:
    """Route root-level archive docs to their extraction targets."""
    stem = Path(filename).stem
    if filename == "Ultimate_Diagnostic.docx":
        records.append(_record("position_spec", "position_audit_v1", build_position_spec(doc), rel))
        return "position_spec"
    if filename.startswith("AI_Provenance_Verification_System_v2.0"):
        records.append(
            _record("provenance_spec", "ai_provenance_v2", build_generic_payload(doc), rel)
        )
        return "provenance_spec"
    if filename.startswith("MODULE_7_AI_Provenance"):
        records.append(
            _record("provenance_spec", "ai_provenance_module7", build_generic_payload(doc), rel)
        )
        return "provenance_spec"
    if filename == "Book_Bible.docx":
        proposals.extend(proposals_from_doc(doc, rel, "INVENTED"))
        return "canon_proposal"
    # Anything else parsed but unmatched: keep it as a generic system doc so
    # coverage stays honest without inventing meaning.
    records.append(_record("system_doc", stem, build_generic_payload(doc), rel))
    return "system_doc"


def _write_coverage_report(data_dir: Path, summary: dict, inventory: list[dict]) -> None:
    out_dir = data_dir / "wa"
    out_dir.mkdir(parents=True, exist_ok=True)
    (out_dir / "coverage_report.json").write_text(
        json.dumps({"summary": summary, "inventory": inventory}, indent=2, ensure_ascii=False),
        encoding="utf-8",
    )
    lines = [
        "# Writing Architect — M0 Decomposition Coverage Report",
        "",
        f"Run `{summary['run_id']}` — {summary['inventory']} archive files, "
        f"{summary['records']} doctrine records, {summary['proposals_seen']} canon proposals "
        f"({summary['proposals_new']} new) in {summary['duration_sec']}s.",
        "",
        "| Status | Count |",
        "|---|---|",
    ]
    for status, n in sorted(summary["coverage"]["by_status"].items()):
        lines.append(f"| {status} | {n} |")
    lines += [
        "",
        "## Every file, accounted for",
        "",
        "| File | Layer | Disposition | Target / Reason |",
        "|---|---|---|---|",
    ]
    for row in sorted(inventory, key=lambda r: r["rel_path"]):
        detail = row.get("target_kind") or row.get("reason") or ""
        if row.get("duplicate_of"):
            detail = f"duplicate of {row['duplicate_of']} — {detail}"
        lines.append(f"| {row['rel_path']} | {row['layer']} | {row['status']} | {detail} |")
    (out_dir / "coverage_report.md").write_text("\n".join(lines) + "\n", encoding="utf-8")
