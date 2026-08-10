"""
Build the Orivellum Installation & Setup Guide DOCX.
Run with: uv run python scripts/build_manual.py
"""

import os

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

os.makedirs("docs/manual", exist_ok=True)
doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

# ── Helpers ───────────────────────────────────────────────────────────────────


def shade_cell(cell, fill_hex: str):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)


def shade_para(p, fill_hex: str):
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    pPr.append(shd)


def h(level: int, text: str, color=None):
    p = doc.add_heading(text, level=level)
    if color:
        for run in p.runs:
            run.font.color.rgb = RGBColor(*color)
    return p


def body(text: str):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    return p


def bullet(text: str):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    return p


def numbered(text: str):
    p = doc.add_paragraph(text, style="List Number")
    p.paragraph_format.space_after = Pt(3)
    return p


def code(text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    shade_para(p, "F0F4F8")
    return p


def note(text: str, fill="FFF8E1", icon="💡", text_rgb=(0x5A, 0x42, 0x00)):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.right_indent = Inches(0.2)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(f"{icon}  {text}")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(*text_rgb)
    shade_para(p, fill)
    return p


def tip(text: str):
    return note(text, "E8F5E9", "✅", (0x1A, 0x5C, 0x2E))


def warn(text: str):
    return note(text, "FDE8E8", "⚠️", (0x8B, 0x1A, 0x1A))


def add_image(path: str, width=6.0, caption=None):
    try:
        doc.add_picture(path, width=Inches(width))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if caption:
            cp = doc.add_paragraph(caption)
            cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cp.paragraph_format.space_after = Pt(12)
            for run in cp.runs:
                run.font.size = Pt(9)
                run.font.italic = True
                run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    except Exception as e:
        body(f"[Image: {caption or path}]  ({e})")


def dark_table_header(row, cols, fill="1a3a2a"):
    for cell in row.cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shade_cell(cell, fill)


def add_table(headers, rows, col_widths, row_fill=("FFFFFF", "F4F7F5")):
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h_text in enumerate(headers):
        tbl.rows[0].cells[i].text = h_text
    dark_table_header(tbl.rows[0], len(headers))
    for idx, data_row in enumerate(rows):
        r = tbl.add_row().cells
        fill = row_fill[idx % 2]
        for j, val in enumerate(data_row):
            r[j].text = str(val)
            shade_cell(r[j], fill)
    for row in tbl.rows:
        for i, cell in enumerate(row.cells):
            if i < len(col_widths):
                cell.width = col_widths[i]
    doc.add_paragraph()
    return tbl


# ─────────────────────────────────────────────────────────────────────────────
# COVER PAGE
# ─────────────────────────────────────────────────────────────────────────────
add_image("docs/manual/hero-banner.jpg", 6.5)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Orivellum")
r.font.size = Pt(36)
r.font.bold = True
r.font.color.rgb = RGBColor(0x1A, 0x3A, 0x2A)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run("Complete Installation & Setup Guide")
r2.font.size = Pt(18)
r2.font.color.rgb = RGBColor(0x9A, 0x7B, 0x2E)

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = p3.add_run("v1.0  ·  August 2026  ·  For: Brian Fochtman")
r3.font.size = Pt(11)
r3.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

doc.add_page_break()

# ─────────────────────────────────────────────────────────────────────────────
# TABLE OF CONTENTS
# ─────────────────────────────────────────────────────────────────────────────
h(1, "Table of Contents")
toc = [
    ("1.", "Quick-Start Checklist"),
    ("2.", "Prerequisites"),
    ("3.", "Installation — Mac / Linux"),
    ("4.", "Installation — Windows"),
    ("5.", "Lemonade AI Server Setup"),
    ("6.", "Recommended LLM Models for Lemonade"),
    ("7.", "First Launch & Configuration"),
    ("8.", "Environment Variables Reference"),
    ("9.", "Troubleshooting"),
]
for num, title in toc:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    nr = p.add_run(f"{num}  ")
    nr.font.bold = True
    nr.font.color.rgb = RGBColor(0x9A, 0x7B, 0x2E)
    p.add_run(title)
doc.add_page_break()

# ─────────────────────────────────────────────────────────────────────────────
# 1. QUICK-START CHECKLIST
# ─────────────────────────────────────────────────────────────────────────────
h(1, "1.  Quick-Start Checklist")
body("Use this checklist before you begin. Check each item off as you complete it.")

add_image("docs/manual/install-steps.jpg", 5.5, "Four-step installation overview")

h(2, "Prerequisites checklist")
for item in [
    "☐  Git 2.35 +",
    "☐  Node.js 18 LTS or 20 LTS",
    "☐  pnpm 8 +  (npm install -g pnpm)",
    "☐  Python 3.12 +",
    "☐  uv  (pip install uv  OR  winget install astral-sh.uv)",
    "☐  8 GB RAM minimum  (16 GB recommended for local LLM)",
    "☐  AMD or NVIDIA GPU  (optional — for Lemonade GPU acceleration)",
]:
    bullet(item)

h(2, "Installation checklist")
for item in [
    "☐  Clone the repository",
    "☐  Run  pnpm install",
    "☐  Run  uv sync",
    "☐  Copy and edit .env  (optional)",
    "☐  Start with  bash scripts/dev.sh",
    "☐  (Optional) Install Lemonade and load a model",
]:
    bullet(item)
doc.add_page_break()

# ─────────────────────────────────────────────────────────────────────────────
# 2. PREREQUISITES
# ─────────────────────────────────────────────────────────────────────────────
h(1, "2.  Prerequisites")

h(2, "Node.js")
body("Download the LTS release from https://nodejs.org. Run the installer and verify with:")
code("node --version   # must be 18.x or 20.x")
tip("Use Node 20 LTS for best compatibility. Versions older than 18 are not supported.")

h(2, "pnpm")
body("Orivellum uses pnpm workspaces. Install it globally after Node.js:")
code("npm install -g pnpm\npnpm --version   # should be 8 or higher")

h(2, "Python 3.12 + and uv")
body("Python 3.12 or newer is required for the API server. uv is the recommended installer:")
code(
    '# Mac / Linux\ncurl -LsSf https://astral.sh/uv/install.sh | sh\n\n# Windows (PowerShell)\npowershell -c "irm https://astral.sh/uv/install.ps1 | iex"'
)
tip(
    "uv is much faster than pip. It creates a virtual environment and pins all dependencies automatically via uv.lock."
)

h(2, "Git")
body("Install from https://git-scm.com or via your package manager:")
code(
    "# Mac\nbrew install git\n\n# Ubuntu / Debian\nsudo apt install git\n\n# Windows — included with Git for Windows installer"
)
doc.add_page_break()

# ─────────────────────────────────────────────────────────────────────────────
# 3. INSTALLATION — MAC / LINUX
# ─────────────────────────────────────────────────────────────────────────────
h(1, "3.  Installation — Mac / Linux")

h(2, "Step 1 — Clone the repository")
code("git clone https://github.com/BFo/orivellum-main.git\ncd orivellum-main")

h(2, "Step 2 — Install JavaScript / TypeScript dependencies")
code("pnpm install")
body(
    "Installs all workspace packages (web UI, shared libraries, API client). "
    "Expect 2–5 minutes on first run."
)

h(2, "Step 3 — Set up the Python virtual environment")
code("uv sync")
body(
    "Creates .venv and installs all Python packages from uv.lock. "
    "Subsequent runs complete in seconds."
)

h(2, "Step 4 — Create your .env file  (optional)")
body("Copy the example and fill in any secrets you need:")
code(
    "cp artifacts/api-server/.env.example artifacts/api-server/.env\n# Edit with nano, vim, or VS Code"
)
tip(
    "The app works without any .env file. "
    "AI model URLs and keys can also be set through the web Settings page."
)

h(2, "Step 5 — Start the application")
code("# API server + Web UI\nbash scripts/dev.sh")
body(
    "The script starts FastAPI on port 8080, waits for a healthy response, then starts Vite on port 5173."
)
body("Open your browser to:")
code("http://localhost:5173")
warn("Keep this terminal open. Press Ctrl+C to stop all services at once.")
doc.add_page_break()

# ─────────────────────────────────────────────────────────────────────────────
# 4. INSTALLATION — WINDOWS
# ─────────────────────────────────────────────────────────────────────────────
h(1, "4.  Installation — Windows")

h(2, "Step 1 — Install prerequisites via winget")
body("Open PowerShell as Administrator and run:")
code(
    "winget install OpenJS.NodeJS.LTS\n"
    "winget install pnpm.pnpm\n"
    "winget install astral-sh.uv\n"
    "winget install Git.Git"
)
body("Restart your terminal after installation so the new tools appear in PATH.")

h(2, "Step 2 — Clone and install")
code(
    "git clone https://github.com/BFo/orivellum-main.git\ncd orivellum-main\npnpm install\nuv sync"
)

h(2, "Step 3 — Development mode  (Vite dev server + API)")
body("In a PowerShell or Git Bash window:")
code(
    "# Git Bash\nbash scripts/dev.sh\n\n# PowerShell — set port then call bash\n$env:API_PORT='8080'; bash scripts/dev.sh"
)
tip(
    "Git for Windows includes Git Bash. If bash is not found, install Git for Windows from https://git-scm.com/download/win"
)

h(2, "Step 4 — Production / Appliance mode  (no dev server)")
body("Builds the UI bundle once, then serves everything from a single FastAPI process:")
code(
    ".\\scripts\\start.ps1\n\n# Skip rebuilding the UI (fast restart):\n.\\scripts\\start.ps1 -SkipBuild"
)
body("Then open:")
code("http://localhost:8080/orivellum-ui/")
tip(
    "Bookmark this URL and use Add to Home Screen in your browser for a near-native app experience."
)

h(2, "Step 5 — Auto-start at Windows login  (optional)")
body("Register Orivellum as a Task Scheduler task so it starts automatically:")
code(".\\scripts\\windows\\register-boot.ps1")
note(
    "After registration Orivellum starts at login. Run register-boot.ps1 -Unregister to remove it."
)
doc.add_page_break()

# ─────────────────────────────────────────────────────────────────────────────
# 5. LEMONADE AI SERVER SETUP
# ─────────────────────────────────────────────────────────────────────────────
h(1, "5.  Lemonade AI Server Setup")
body(
    "Lemonade Server is AMD's official local LLM inference tool for Ryzen AI hardware. "
    "Unlike generic inference tools, Lemonade automatically schedules workloads across "
    "the NPU (XDNA2), iGPU (RDNA 3.5), and CPU — no manual GPU flags needed. "
    "Every conversation and document stays 100% private and offline."
)

add_image(
    "docs/manual/lemonade-setup.jpg",
    5.5,
    "Lemonade auto-schedules across NPU, iGPU, and CPU — no manual GPU config required",
)

h(2, "Install Lemonade")
body("Download the Windows installer from the official site:")
code(
    "# 1. Visit https://lemonade-server.ai and download the Windows installer\n"
    "# 2. Run the installer\n"
    "# 3. Lemonade registers itself to START AUTOMATICALLY AT LOGIN — no manual launch needed\n\n"
    "# Verify it is running:\n"
    "Invoke-WebRequest http://127.0.0.1:13305/api/v1/models -UseBasicParsing\n"
    '# Expected: {"status":"ok"}'
)
tip(
    "Lemonade installs as a background service. After the first install it starts automatically — "
    "you will never need to launch it manually again."
)

h(2, "Pull your models")
body(
    "Use lemonade pull to download models. Model names must match the Lemonade catalog "
    "EXACTLY — they double as the API model IDs. These are the models Orivellum is "
    "configured to use:"
)
code(
    "# Workhorse — daily driver, MoE (~3 B active), vision built in (~23 GB)\n"
    "lemonade pull Qwen3.6-35B-A3B-GGUF\n\n"
    "# Reasoner — 120 B MoE in native MXFP4 (~63 GB), best local reasoning\n"
    "lemonade pull gpt-oss-120b-mxfp-GGUF\n\n"
    "# Coder — agentic coding, 256 K context (~19 GB)\n"
    "lemonade pull Qwen3-Coder-30B-A3B-Instruct-GGUF\n\n"
    "# Embeddings for semantic search (~8 GB)\n"
    "lemonade pull Qwen3-Embedding-8B-GGUF\n\n"
    "# OPTIONAL: fast low-power NPU model for quick Q&A\n"
    "lemonade pull gpt-oss-20b-NPU\n\n"
    "# Check what you have:\n"
    "lemonade list"
)
tip(
    "Why no dense 70B? This machine's unified memory moves ~256 GB/s. A dense 70 B model "
    "reads all ~40 GB of weights for EVERY token (~4-5 tokens/sec). MoE models like "
    "gpt-oss-120b activate only ~5 B parameters per token — 30-40 tokens/sec at higher "
    "quality. On this hardware, MoE wins every time."
)

h(2, "Connect Orivellum to Lemonade")
numbered("Open Orivellum in your browser")
numbered("Go to System page → Settings section")
numbered('Under "Local AI (Lemonade)" verify the URL is set to:')
code("http://127.0.0.1:13305/api/v1")
numbered("The model slots are already configured in config.yaml — no manual model ID entry needed")
numbered("Click Save Settings — the green indicator confirms the connection")

h(2, "Hardware scheduling (automatic)")
body("Lemonade picks the best hardware for each model size automatically:")
bullet("NPU (XDNA2, 50 TOPS): 7–20 B models — low power, sub-second first token (gpt-oss-20b-NPU)")
bullet("iGPU (40 RDNA 3.5 CUs, ≈ RTX 4070): the big models — Qwen3.6-35B-A3B, gpt-oss-120b")
bullet("Unified 128 GB LPDDR5X: no separate VRAM — even the 120 B reasoner fits entirely on-chip")
tip(
    "The Ryzen AI Max+ 395 has ~112 GB allocatable to models. The workhorse (~23 GB) and "
    "embedder (~8 GB) stay resident; the reasoner (~63 GB) and coder (~19 GB) load on "
    "demand — everything fits."
)
doc.add_page_break()

# ─────────────────────────────────────────────────────────────────────────────
# 6. RECOMMENDED LLM MODELS
# ─────────────────────────────────────────────────────────────────────────────
h(1, "6.  Orivellum Model Slots & Recommended Models")
body(
    "Orivellum uses six named model slots, each optimised for a specific type of work. "
    "All slots are configured in config.yaml in the project root — you do not set them "
    "one by one in the UI. The defaults below are tuned for the AMD Ryzen AI Max+ 395."
)

slots_headers = ["Slot", "config.yaml key", "Purpose", "Recommended model", "Hardware used"]
slots_rows = [
    (
        "Workhorse",
        "workhorse_model",
        "Chat, summarisation, general\nreasoning, web research",
        "Qwen3.6-35B-A3B-GGUF\n(MoE, 3 B active — 50-70 tok/s)",
        "iGPU (RDNA 3.5)\n~23 GB of 112 GB",
    ),
    (
        "Reasoner",
        "reasoner_model",
        "Complex multi-step reasoning,\nverification, deep analysis",
        "gpt-oss-120b-mxfp-GGUF\n(MoE, 5.1 B active — 30-40 tok/s)",
        "iGPU (RDNA 3.5)\n~63 GB of 112 GB",
    ),
    (
        "Coder",
        "coder_model",
        "Code generation, document\nworkshop, structured output",
        "Qwen3-Coder-30B-A3B-Instruct-GGUF\n(alt: Qwen3-Coder-Next-GGUF ~48 GB)",
        "iGPU\n~19 GB",
    ),
    (
        "Embedder",
        "embedder_model",
        "Semantic search vectors for\nall documents and knowledge",
        "Qwen3-Embedding-8B-GGUF\n(alt: nomic-embed-text-v2-moe-GGUF)",
        "iGPU / CPU\n~8 GB",
    ),
    (
        "Vision",
        "vision_model",
        "OCR of scanned PDFs and\nimages (Tesseract fallback)",
        "Qwen3.6-35B-A3B-GGUF\n(same as Workhorse — no extra download)",
        "iGPU",
    ),
    (
        "Reranker",
        "reranker_model",
        "Re-ranks search results for\nbetter retrieval quality",
        "bge-reranker-v2-m3-GGUF\n(optional — leave empty)",
        "CPU / NPU",
    ),
]
add_table(
    slots_headers,
    slots_rows,
    [Inches(0.75), Inches(1.05), Inches(1.4), Inches(1.65), Inches(1.0)],
)
tip(
    "The Ryzen AI Max+ 395 has 128 GB unified memory with ~112 GB allocatable. "
    "The workhorse (~23 GB) and embedder (~8 GB) stay resident; the reasoner (~63 GB) "
    "and coder (~19 GB) load on demand — no CPU offload, no paging, full GPU speed."
)
note(
    "gpt-oss-120b is OpenAI's open-weight 120 B Mixture-of-Experts model, shipped in "
    "native MXFP4. Only ~5.1 B parameters are active per token, so it is BOTH the best "
    "local reasoner available AND fast on this machine (~30-40 tokens/sec)."
)

h(2, "Pulling all models at once")
body(
    "Run these commands once after Lemonade is installed — they download everything Orivellum needs. "
    "Names must match the Lemonade catalog exactly (they double as API model IDs):"
)
code(
    "lemonade pull Qwen3.6-35B-A3B-GGUF                # Workhorse — MoE, vision built in\n"
    "lemonade pull gpt-oss-120b-mxfp-GGUF              # Reasoner  — best local reasoning\n"
    "lemonade pull Qwen3-Coder-30B-A3B-Instruct-GGUF   # Coder     — 256 K context\n"
    "lemonade pull Qwen3-Embedding-8B-GGUF             # Embedder  — semantic search\n"
    "lemonade pull gpt-oss-20b-NPU                     # Optional  — fast NPU model\n"
    "lemonade pull bge-reranker-v2-m3-GGUF             # Optional  — search reranker\n\n"
    "# Verify everything downloaded:\n"
    "lemonade list"
)

h(2, "config.yaml — default model configuration")
body(
    "The project ships with these defaults already set in config.yaml. "
    "Edit the file to swap models at any time — no restart required for most settings:"
)
code(
    "serving:\n"
    '  workhorse_model: "Qwen3.6-35B-A3B-GGUF"              # General chat & reasoning\n'
    '  reasoner_model:  "gpt-oss-120b-mxfp-GGUF"            # Deep analysis (brain icon in chat)\n'
    '  coder_model:     "Qwen3-Coder-30B-A3B-Instruct-GGUF" # Code & structured output\n'
    '  embedder_model:  "Qwen3-Embedding-8B-GGUF"           # Semantic search\n'
    '  vision_model:    "Qwen3.6-35B-A3B-GGUF"              # OCR (same model as workhorse)\n'
    '  reranker_model:  ""                                   # Leave empty → no reranking'
)
tip(
    "In the chat interface, tap the brain icon (🧠) to switch from the Workhorse model "
    "to the Reasoner (gpt-oss-120b) for questions that need deeper analysis."
)
doc.add_page_break()

# ─────────────────────────────────────────────────────────────────────────────
# 7. FIRST LAUNCH & CONFIGURATION
# ─────────────────────────────────────────────────────────────────────────────
h(1, "7.  First Launch & Configuration")

h(2, "What you see on first launch")
numbered("The web UI opens at  http://localhost:5173")
numbered(
    "The Home Screen shows a launch tile for each app: Writing, Learning, Chat, Studio, Command, Mail, and Library"
)
numbered("The Library app is where you upload your first documents")
numbered("Chat works immediately — try asking anything even without a local model")

h(2, "Connect a local AI model")
numbered("Start Lemonade and load a model  (see Section 5)")
numbered("Open Orivellum → Command app → System page → Settings card")
numbered('Under "Local AI (Lemonade)" enter:  http://127.0.0.1:13305/api/v1')
numbered("Click Save — the status indicator turns green")

h(2, "Import your first document")
numbered("Open the Library app from the Home Screen")
numbered("Drag and drop a PDF, DOCX, EPUB, or image — or click to browse")
numbered("Orivellum extracts knowledge, auto-assigns it to a Work, and makes it searchable in Chat")
tip(
    "Works are like projects or books — they group related documents. "
    "Orivellum automatically suggests the right Work when you import."
)

h(2, "Start your first chat")
body(
    "Open the Chat app from the Home Screen. "
    "Scope a conversation to a specific Work using the Work selector at the top — "
    "the AI will draw exclusively from that Work's documents for precise, cited answers."
)
doc.add_page_break()

# ─────────────────────────────────────────────────────────────────────────────
# 8. ENVIRONMENT VARIABLES REFERENCE
# ─────────────────────────────────────────────────────────────────────────────
h(1, "8.  Environment Variables Reference")
body(
    "Create  artifacts/api-server/.env  (copy from .env.example). All variables are optional — sensible defaults apply."
)

env_headers = ["Variable", "Default", "Description"]
env_rows = [
    ("API_PORT", "8080", "Port for the FastAPI server"),
    (
        "SESSION_SECRET",
        "(auto-generated)",
        "Cookie session secret — set a strong random string in production",
    ),
    ("LEMONADE_URL", "http://127.0.0.1:13305/api/v1", "Lemonade API base URL"),
    ("LEMONADE_MODEL", "(blank = server default)", "Model ID to request from Lemonade"),
    ("AI_EXTRACTION_ENABLED", "false", "Enable LLM-powered knowledge extraction on import"),
    ("TAVILY_API_KEY", "(blank)", "Tavily search key for online research feature"),
    ("DATA_DIR", "./data", "Directory where Orivellum stores its database and files"),
    ("DEBUG", "0", "Set to 1 for verbose API logging"),
]
add_table(env_headers, env_rows, [Inches(1.9), Inches(1.6), Inches(2.5)])
doc.add_page_break()

# ─────────────────────────────────────────────────────────────────────────────
# 9. TROUBLESHOOTING
# ─────────────────────────────────────────────────────────────────────────────
h(1, "9.  Troubleshooting")

h(2, "API server does not start")
bullet("Check Python version:  python --version  — must be 3.12 or newer")
bullet("Re-run  uv sync  to ensure all packages are installed")
bullet(
    "Check port 8080 is free:  lsof -i :8080  (Mac/Linux)  or  netstat -ano | findstr :8080  (Windows)"
)
code("# Start with verbose output to see the error\nDEBUG=1 uv run python -m orivellum.api.main")

h(2, "Web UI shows a blank screen")
bullet("Ensure the API server started — dev.sh waits for it but check for crash messages")
bullet("Open browser console (F12) and look for red error messages")
bullet("Hard-reload:  Ctrl+Shift+R  (Windows/Linux)  or  Cmd+Shift+R  (Mac)")

h(2, "Lemonade connection fails in Settings")
bullet("Verify Lemonade is running:  curl http://127.0.0.1:13305/api/v1/models")
bullet("Check the URL exactly matches the Settings field — no trailing slash")
bullet("Ensure a model is loaded before sending messages")
code(
    "# Test Lemonade from the terminal\ncurl http://127.0.0.1:13305/api/v1/models\n\n"
    '# Expected response:\n# {"object":"list","data":[{"id":"..."}]}'
)

h(2, "Slow AI responses with local model")
bullet("Switch to a smaller quantization: Q4_K_M instead of Q8_0")
bullet("Use a smaller model: Phi-4 Mini 3.8B instead of Llama 3.1 8B")
bullet("Ensure Lemonade is using the GPU, not CPU:  lemonade status")
tip("Q4_K_M at 8B offers 90% of Q8_0 quality at about 55% of the file size and memory usage.")

h(2, "pnpm install fails on Windows")
bullet("Run PowerShell as Administrator")
bullet("Set execution policy:  Set-ExecutionPolicy RemoteSigned")
bullet("Then re-run:  pnpm install")

# ─────────────────────────────────────────────────────────────────────────────
# Save
# ─────────────────────────────────────────────────────────────────────────────
out = "docs/manual/Orivellum_Installation_Guide.docx"
doc.save(out)
print(f"Saved: {out}")
