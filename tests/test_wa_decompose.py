"""Tests for the Writing Architect archive decomposer (Pipeline M0)."""

from __future__ import annotations

import io
import zipfile
from pathlib import Path

import pytest

from orivellum.capabilities.wa_decompose import (
    _dedupe_key,
    layer_of,
    parse_markdown,
    run_decompose,
)
from orivellum.database.db import OrivellumDB
from orivellum.database.wa_store import WAStore

WS = "WRITING_ARCHITECT/Module writing/WRITING_SYSTEM"


def _docx_bytes(title: str, sections: dict[str, list[str]]) -> bytes:
    import docx

    d = docx.Document()
    d.add_paragraph(title)
    for heading, items in sections.items():
        d.add_heading(heading, level=1)
        for item in items:
            d.add_paragraph(item, style="List Bullet")
    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


def _index_docx(rows: list[list[str]]) -> bytes:
    import docx

    d = docx.Document()
    d.add_paragraph("ENGINE_INDEX")
    table = d.add_table(rows=0, cols=4)
    for row in [["File", "Purpose", "When to Call", "Runtime Status"], *rows]:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val
    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


BIBLE_MD = """---
title: 'Iron Age I-II Material Culture'
note_type: research
---

# Iron Age Material Culture

## Architecture

The pillared house was the dominant domestic form of Iron Age Israel,
divided into four rooms by rough limestone pillars shared with animals.

## Clothing

Wool tunics were the common garment; linen marked wealth and priesthood,
and dyed cloth signified rank throughout the ancient Near East.
"""


@pytest.fixture()
def archive(tmp_path: Path) -> Path:
    engine_doc = _docx_bytes(
        "ENGINE_01_TEST",
        {
            "Purpose": ["Test consequence modeling"],
            "Required Inputs": ["Baseline", "Trigger"],
            "Allowed Operations": ["Build", "Review"],
            "Forbidden Operations": ["Lock"],
            "Stop Conditions": ["Baseline underbuilt"],
        },
    )
    variant_a = _docx_bytes("VARIANT", {"Purpose": ["First version of the doctrine"]})
    variant_b = _docx_bytes("VARIANT", {"Purpose": ["Second, different version entirely"]})
    zpath = tmp_path / "wa.zip"
    with zipfile.ZipFile(zpath, "w") as zf:
        zf.writestr(f"{WS}/ENGINE/ENGINE_01_TEST.docx", engine_doc)
        zf.writestr(
            f"{WS}/ENGINE/ENGINE_INDEX.docx",
            _index_docx([["ENGINE_01_TEST", "Testing", "When testing", "On-Demand"]]),
        )
        # Identical duplicate pair → deduped
        zf.writestr(f"{WS}/ENGINE/ENGINE_01_TEST__1.docx", engine_doc)
        # Differing pair → canonical extracted, variant deferred
        zf.writestr(f"{WS}/RUNTIME/runtime_test_policy.docx", variant_a)
        zf.writestr(f"{WS}/RUNTIME/runtime_test_policy__1.docx", variant_b)
        zf.writestr(
            f"{WS}/RUNTIME/schema_test_record.docx",
            _docx_bytes("schema_test_record", {"Fields": ["id", "status"]}),
        )
        zf.writestr(
            f"{WS}/STORY_SYSTEM/CHAPTER_TEST_SCHEMA.docx",
            _docx_bytes("CHAPTER_TEST_SCHEMA", {"Fields": ["chapter_id"]}),
        )
        zf.writestr(f"{WS}/BIBLE_DATA/Iron_Age_Material_Culture.md", BIBLE_MD)
        zf.writestr(
            "WRITING_ARCHITECT/Ultimate_Diagnostic.docx",
            _docx_bytes(
                "Diagnostic",
                {
                    "Audit": [
                        "1. PLOT HOLES",
                        "Look at the timeline versus the settings for contradictions.",
                        "2. PACING",
                        "Pinpoint where narrative momentum is at risk of stalling.",
                    ]
                },
            ),
        )
        zf.writestr(
            "WRITING_ARCHITECT/ULTIMATE_PROSE_SYSTEM_v2.1_IMPLEMENTATION.md",
            "# Prose\n\n## Rhythm\n- Never open consecutive paragraphs with the same word\n"
            "- No more than 2 semicolons per chapter\n",
        )
        # Nested archive and resource forks must be handled
        zf.writestr("WRITING_ARCHITECT/nested_repo.zip", b"PK\x05\x06" + b"\x00" * 18)
        zf.writestr("__MACOSX/WRITING_ARCHITECT/._junk", b"\x00")
        zf.writestr(f"{WS}/ENGINE/._ENGINE_01_TEST.docx", b"\x00")
    return zpath


@pytest.fixture()
def db(tmp_path: Path) -> OrivellumDB:
    return OrivellumDB(str(tmp_path / "test.db"))


def test_dedupe_key_variants():
    assert _dedupe_key("a/ENGINE_INDEX__1.docx") == "a/ENGINE_INDEX.docx"
    assert _dedupe_key("a/ENGINE_INDEX__2.docx") == "a/ENGINE_INDEX.docx"
    assert _dedupe_key("a/SOVEREIGN_MASTER_v1.4 2") == "a/SOVEREIGN_MASTER_v1.4"
    assert _dedupe_key("a/plain.docx") == "a/plain.docx"


def test_layer_routing():
    assert layer_of(f"{WS}/ENGINE/x.docx") == "ENGINE"
    assert layer_of(f"{WS}/BIBLE_DATA/x.md") == "BIBLE_DATA"
    assert layer_of("WRITING_ARCHITECT/Book_Bible.docx") == "ROOT"


def test_markdown_frontmatter():
    doc = parse_markdown(BIBLE_MD, "fallback")
    assert doc.frontmatter["note_type"] == "research"
    assert any(s.heading == "Architecture" for s in doc.sections)


def test_full_decompose_coverage(archive: Path, db: OrivellumDB, tmp_path: Path):
    summary = run_decompose(archive, db, data_dir=tmp_path / "data")
    cov = summary["coverage"]
    assert cov["fully_accounted"] is True
    # 11 real files (2 resource-fork/__MACOSX entries excluded)
    assert cov["total_docs"] == 11
    store = WAStore(db)

    # Identical duplicate → deduped with pointer to canonical
    inv = {r["rel_path"]: r for r in store.list_inventory()}
    dup = inv[f"{WS}/ENGINE/ENGINE_01_TEST__1.docx"]
    assert dup["status"] == "deduped"
    assert dup["duplicate_of"] == f"{WS}/ENGINE/ENGINE_01_TEST.docx"

    # Differing variant → explicitly deferred, canonical extracted
    variant_rows = [r for r in inv.values() if "runtime_test_policy" in r["rel_path"]]
    statuses = sorted(r["status"] for r in variant_rows)
    assert statuses == ["deferred", "extracted"]

    # Nested zip deferred with a reason, never silently dropped
    nested = inv["WRITING_ARCHITECT/nested_repo.zip"]
    assert nested["status"] == "deferred" and nested["reason"]

    # Coverage report files written
    assert (tmp_path / "data/wa/coverage_report.md").exists()
    assert (tmp_path / "data/wa/coverage_report.json").exists()


def test_engine_contract_fields(archive: Path, db: OrivellumDB, tmp_path: Path):
    run_decompose(archive, db, data_dir=tmp_path / "data")
    store = WAStore(db)
    contracts = store.list_records(record_type="engine_contract")
    rec = store.get_record(next(r["id"] for r in contracts if r["name"] == "ENGINE_01_TEST"))
    p = rec["payload"]
    assert p["purpose"] == ["Test consequence modeling"]
    assert p["allowed_operations"] == ["Build", "Review"]
    assert p["forbidden_operations"] == ["Lock"]
    # Backfilled from the ENGINE_INDEX operator table
    assert p["certification_status"] == "On-Demand"
    assert p["when_to_call"] == "When testing"


def test_record_type_routing(archive: Path, db: OrivellumDB, tmp_path: Path):
    run_decompose(archive, db, data_dir=tmp_path / "data")
    store = WAStore(db)
    types = {r["name"]: r["record_type"] for r in store.list_records()}
    assert types["schema_test_record"] == "record_schema"
    # The differing pair yields exactly one canonical runtime policy
    runtime_names = [n for n, t in types.items() if n.startswith("runtime_test_policy")]
    assert len(runtime_names) == 1
    assert types[runtime_names[0]] == "runtime_policy"
    assert types["CHAPTER_TEST_SCHEMA"] == "table_spec"
    assert types["position_audit_v1"] == "position_spec"
    assert types["voice_envelope_v1"] == "voice_envelope"


def test_position_spec_dimensions(archive: Path, db: OrivellumDB, tmp_path: Path):
    run_decompose(archive, db, data_dir=tmp_path / "data")
    store = WAStore(db)
    rec = store.get_record(next(r["id"] for r in store.list_records(record_type="position_spec")))
    dims = rec["payload"]["dimensions"]
    assert [d["dimension"] for d in dims] == ["PLOT HOLES", "PACING"]
    assert dims[0]["probes"]


def test_voice_envelope_measurable_rules(archive: Path, db: OrivellumDB, tmp_path: Path):
    run_decompose(archive, db, data_dir=tmp_path / "data")
    store = WAStore(db)
    rec = store.get_record(next(r["id"] for r in store.list_records(record_type="voice_envelope")))
    rules = rec["payload"]["measurable_constraints"]["ULTIMATE_PROSE_SYSTEM_v2.1_IMPLEMENTATION.md"]
    assert any("semicolons" in r["rule"] for r in rules)


def test_canon_proposals_never_auto_accepted(archive: Path, db: OrivellumDB, tmp_path: Path):
    run_decompose(archive, db, data_dir=tmp_path / "data")
    store = WAStore(db)
    props = store.list_proposals()
    assert props, "BIBLE_DATA must yield proposals"
    assert all(p["status"] == "proposed" for p in props)
    # research frontmatter → HISTORICAL default
    assert all(p["classification"] == "HISTORICAL" for p in props)
    assert all(p["scope"] == "series:The Harp in the Dark" for p in props)
    assert all(p["source_path"] and p["source_location"] for p in props)


def test_rerun_preserves_ratification(archive: Path, db: OrivellumDB, tmp_path: Path):
    run_decompose(archive, db, data_dir=tmp_path / "data")
    store = WAStore(db)
    target = store.list_proposals()[0]
    decided = store.decide_proposal(target["id"], "approved")
    assert decided["status"] == "approved" and decided["decided_at"]

    summary2 = run_decompose(archive, db, data_dir=tmp_path / "data")
    assert summary2["proposals_new"] == 0
    again = [p for p in store.list_proposals() if p["id"] == target["id"]][0]
    assert again["status"] == "approved"


def test_decide_proposal_guards(db: OrivellumDB):
    store = WAStore(db)
    assert store.decide_proposal("missing-id", "approved") is None
    with pytest.raises(ValueError):
        store.decide_proposal("missing-id", "bogus")
